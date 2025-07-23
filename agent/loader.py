# agent/loader.py
import fitz  # PyMuPDF
from bs4 import BeautifulSoup
import requests
from readability.readability import Document
from docx import Document as DocxDocument


def load_pdf(path):
    doc = fitz.open(path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text


def load_html(path_or_url):
    headers = {"User-Agent": "Mozilla/5.0"}
    if path_or_url.startswith("http"):
        response = requests.get(path_or_url, headers=headers, timeout=10)
        response.raise_for_status()
        doc = Document(response.text)
        readable_html = doc.summary()
    else:
        with open(path_or_url, "r", encoding="utf-8") as f:
            readable_html = f.read()

    soup = BeautifulSoup(readable_html, "html.parser")
    texts = []
    for tag in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        txt = tag.get_text(strip=True)
        if txt and len(txt) > 30:
            texts.append(txt)
    return "\n".join(texts)


def extract_seo_signals(url: str) -> dict:
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        r = requests.get(url, headers=headers, timeout=10)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")

        title = soup.title.string.strip() if soup.title else ""
        meta_desc = ""
        if soup.find("meta", attrs={"name": "description"}):
            meta_desc = soup.find("meta", attrs={"name": "description"}).get("content", "").strip()

        headings = [h.get_text(strip=True) for h in soup.find_all(["h1", "h2", "h3"])]
        links = [a["href"] for a in soup.find_all("a", href=True)]
        cta_count = len([a for a in links if any(k in a.lower() for k in ["kontakt", "buchen", "jetzt", "termin"])])

        return {
            "title": title,
            "meta_description": meta_desc,
            "headings": headings[:10],
            "num_links": len(links),
            "cta_links": cta_count
        }
    except Exception as e:
        return {"error": str(e)}


def load_docx(path):
    doc = DocxDocument(path)
    text = "\n".join(para.text for para in doc.paragraphs)
    return text
